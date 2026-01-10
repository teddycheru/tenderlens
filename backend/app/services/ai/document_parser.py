"""
Document parser for extracting text from PDF and DOCX files.
"""

from typing import Optional
import io
import re


class DocumentParser:
    """
    Extract text from PDF and DOC files.

    Note: Requires PyPDF2 and python-docx to be installed.
    Install with: pip install PyPDF2 python-docx
    """

    @staticmethod
    def extract_from_pdf(file_content: bytes) -> str:
        """
        Extract text from PDF file.

        Args:
            file_content: Binary content of PDF file

        Returns:
            Extracted text as string

        Raises:
            Exception: If PDF parsing fails
            ImportError: If PyPDF2 is not installed
        """
        try:
            import PyPDF2
        except ImportError:
            raise ImportError(
                "PyPDF2 is required for PDF parsing. "
                "Install with: pip install PyPDF2==3.0.1"
            )

        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)

            text = ""
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

            return DocumentParser.clean_text(text)
        except Exception as e:
            raise Exception(f"PDF parsing error: {str(e)}")

    @staticmethod
    def extract_from_docx(file_content: bytes) -> str:
        """
        Extract text from DOCX file.

        Args:
            file_content: Binary content of DOCX file

        Returns:
            Extracted text as string

        Raises:
            Exception: If DOCX parsing fails
            ImportError: If python-docx is not installed
        """
        try:
            import docx
        except ImportError:
            raise ImportError(
                "python-docx is required for DOCX parsing. "
                "Install with: pip install python-docx==1.1.0"
            )

        try:
            doc_file = io.BytesIO(file_content)
            doc = docx.Document(doc_file)

            text = ""
            for paragraph in doc.paragraphs:
                if paragraph.text:
                    text += paragraph.text + "\n"

            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text:
                            text += cell.text + " "
                    text += "\n"

            return DocumentParser.clean_text(text)
        except Exception as e:
            raise Exception(f"DOCX parsing error: {str(e)}")

    @staticmethod
    def clean_text(text: str) -> str:
        """
        Clean extracted text by removing excessive whitespace and special characters.

        Args:
            text: Raw extracted text

        Returns:
            Cleaned text
        """
        if not text:
            return ""

        # Remove excessive whitespace (multiple spaces, tabs, newlines)
        text = re.sub(r'\s+', ' ', text)

        # Remove special control characters but keep punctuation
        text = re.sub(r'[^\w\s.,;:!?()\-\'\"@#+=/\[\]{}]+', '', text)

        # Remove leading/trailing whitespace
        text = text.strip()

        return text

    @staticmethod
    def extract_text(file_content: bytes, filename: str) -> str:
        """
        Extract text based on file type (auto-detect from filename).

        Args:
            file_content: Binary content of file
            filename: Name of file (used to determine type)

        Returns:
            Extracted text as string

        Raises:
            ValueError: If file type is not supported
            Exception: If parsing fails
        """
        filename_lower = filename.lower()

        if filename_lower.endswith('.pdf'):
            return DocumentParser.extract_from_pdf(file_content)
        elif filename_lower.endswith('.docx') or filename_lower.endswith('.doc'):
            # Note: .doc (old format) requires additional library (pywin32 on Windows)
            # We'll attempt to parse as .docx
            return DocumentParser.extract_from_docx(file_content)
        else:
            raise ValueError(
                f"Unsupported file type: {filename}. "
                "Only PDF and DOCX files are supported."
            )

    @staticmethod
    def get_word_count(text: str) -> int:
        """
        Count words in text.

        Args:
            text: Text to count words in

        Returns:
            Number of words
        """
        if not text:
            return 0
        return len(text.split())

    @staticmethod
    def get_char_count(text: str) -> int:
        """
        Count characters in text (excluding whitespace).

        Args:
            text: Text to count characters in

        Returns:
            Number of characters
        """
        if not text:
            return 0
        return len(text.replace(" ", "").replace("\n", "").replace("\t", ""))


# Global document parser instance
document_parser = DocumentParser()
